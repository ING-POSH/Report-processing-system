#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Document Processor Module
Handles extraction of text from various document formats
Authors: AI Assistant
Date: February 26, 2026
"""

import os
import logging
from typing import Optional
from abc import ABC, abstractmethod

try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import python_magic
    HAS_MAGIC = True
except ImportError:
    HAS_MAGIC = False

logger = logging.getLogger(__name__)

class DocumentProcessorInterface(ABC):
    """Abstract base class for document processors"""
    
    @abstractmethod
    def can_process(self, file_path: str) -> bool:
        """Check if this processor can handle the file"""
        pass
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """Extract text from the document"""
        pass

class WordDocumentProcessor(DocumentProcessorInterface):
    """Processor for Microsoft Word documents (.docx, .doc)"""
    
    def __init__(self):
        if not HAS_PYTHON_DOCX:
            raise ImportError("python-docx library is required for Word processing")
    
    def can_process(self, file_path: str) -> bool:
        """Check if file is a Word document"""
        extension = os.path.splitext(file_path)[1].lower()
        return extension in ['.docx', '.doc']
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {e}")
            raise

class PDFDocumentProcessor(DocumentProcessorInterface):
    """Processor for PDF documents"""
    
    def __init__(self):
        if not HAS_PDFPLUMBER and not HAS_PYPDF2:
            raise ImportError("Either pdfplumber or PyPDF2 is required for PDF processing")
    
    def can_process(self, file_path: str) -> bool:
        """Check if file is a PDF"""
        return os.path.splitext(file_path)[1].lower() == '.pdf'
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from PDF using the best available method"""
        try:
            # Try pdfplumber first (better quality)
            if HAS_PDFPLUMBER:
                return self._extract_with_pdfplumber(file_path)
            # Fallback to PyPDF2
            elif HAS_PYPDF2:
                return self._extract_with_pypdf2(file_path)
            else:
                raise ImportError("No PDF processing library available")
                
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise
    
    def _extract_with_pdfplumber(self, file_path: str) -> str:
        """Extract text using pdfplumber"""
        text_parts = []
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text.strip())
        
        return '\n\n'.join(text_parts)
    
    def _extract_with_pypdf2(self, file_path: str) -> str:
        """Extract text using PyPDF2 (fallback)"""
        text_parts = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text.strip())
        
        return '\n\n'.join(text_parts)

class TextDocumentProcessor(DocumentProcessorInterface):
    """Processor for plain text files"""
    
    def can_process(self, file_path: str) -> bool:
        """Check if file is a text file"""
        extension = os.path.splitext(file_path)[1].lower()
        text_extensions = ['.txt', '.md', '.rtf', '.log']
        return extension in text_extensions
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        if content.strip():
                            return content.strip()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read().strip()
                
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            raise

class DocumentProcessor:
    """Main document processor that delegates to appropriate handlers"""
    
    def __init__(self, config):
        self.config = config
        self.processors = []
        
        # Initialize available processors
        try:
            self.processors.append(WordDocumentProcessor())
            logger.info("Word document processor initialized")
        except ImportError as e:
            logger.warning(f"Word processor not available: {e}")
        
        try:
            self.processors.append(PDFDocumentProcessor())
            logger.info("PDF document processor initialized")
        except ImportError as e:
            logger.warning(f"PDF processor not available: {e}")
        
        self.processors.append(TextDocumentProcessor())
        logger.info("Text document processor initialized")
        
        logger.info(f"Document processor initialized with {len(self.processors)} processors")
    
    def get_file_type(self, file_path: str) -> str:
        """Determine file type using magic numbers or extension"""
        if HAS_MAGIC:
            try:
                mime_type = python_magic.from_file(file_path, mime=True)
                if mime_type.startswith('application/vnd.openxmlformats-officedocument'):
                    return 'word'
                elif mime_type == 'application/pdf':
                    return 'pdf'
                elif mime_type.startswith('text/'):
                    return 'text'
            except Exception:
                pass
        
        # Fallback to extension
        extension = os.path.splitext(file_path)[1].lower()
        if extension in ['.docx', '.doc']:
            return 'word'
        elif extension == '.pdf':
            return 'pdf'
        elif extension in ['.txt', '.md', '.rtf']:
            return 'text'
        else:
            return 'unknown'
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from document using appropriate processor"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Find appropriate processor
        for processor in self.processors:
            if processor.can_process(file_path):
                logger.info(f"Processing {file_path} with {processor.__class__.__name__}")
                return processor.extract_text(file_path)
        
        # No processor found
        file_type = self.get_file_type(file_path)
        raise ValueError(f"No processor available for file type: {file_type} ({file_path})")
    
    def get_supported_formats(self) -> list:
        """Get list of supported file formats"""
        return self.config.SUPPORTED_FORMATS['documents']

# Example usage
if __name__ == "__main__":
    # Test the document processor
    config = type('Config', (), {
        'SUPPORTED_FORMATS': {
            'documents': ['.docx', '.doc', '.pdf', '.txt', '.md']
        }
    })
    
    processor = DocumentProcessor(config)
    
    # Test with sample files (would need actual files to test)
    print("Supported document formats:", processor.get_supported_formats())